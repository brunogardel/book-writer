"""
Importar — upload de capítulos individuais ou livro completo.
Suporta .txt, .md, .docx e .pdf.
"""

import sys
import re
import uuid
from pathlib import Path
from io import BytesIO

ROOT = Path(__file__).parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

import streamlit as st
from utils.storage import init_session_state, get_chapters, save_chapter, load_book, save_book

st.set_page_config(page_title="Importar — Escritor", page_icon="📥", layout="wide")
init_session_state()

st.title("📥 Importar Texto")
st.markdown(
    "Importe capítulos ou o livro inteiro a partir de arquivos. "
    "Formatos suportados: **TXT**, **Markdown (.md)**, **Word (.docx)**, **PDF**."
)


# ─── Funções de extração de texto ────────────────────────────────────────────

def extract_txt(file_bytes: bytes) -> str:
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(enc)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def extract_docx(file_bytes: bytes) -> str:
    try:
        import docx
        doc = docx.Document(BytesIO(file_bytes))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except ImportError:
        st.error("Biblioteca `python-docx` não encontrada. Execute: `pip install python-docx`")
        return ""
    except Exception as e:
        st.error(f"Erro ao ler DOCX: {e}")
        return ""


def _docx_collect_paragraphs(file_bytes: bytes) -> list[dict] | None:
    """Coleta metadados de todos os parágrafos de um .docx."""
    SKIP_TEXTS = {"bottom of form", "top of form"}
    try:
        import docx
        doc = docx.Document(BytesIO(file_bytes))
    except Exception:
        return None

    result = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or text.lower() in SKIP_TEXTS:
            continue
        style = para.style.name if para.style else "Normal"
        runs = [r for r in para.runs if r.text.strip()]
        # bold=True: explícito; bold=None: herdado do estilo (tratamos como bold se nenhum run
        # diz explicitamente False — cobre títulos com estilo bold aplicado via tema do Word)
        is_bold = bool(runs) and not any(r.bold is False for r in runs)
        result.append({
            "text": text,
            "style": style,
            "style_lower": style.lower(),
            "is_bold": is_bold,
            "is_short": len(text) <= 80,
        })
    return result


def split_docx_into_chapters(file_bytes: bytes) -> tuple[list[dict] | None, str]:
    """
    Divide .docx em capítulos usando 3 estratégias em cascata:
      1. Estilos padrão Word (Heading 1/2/3, Title, Título 1/2/3…)
      2. Estilo customizado dominante entre parágrafos curtos
      3. Parágrafos curtos em negrito

    Retorna (capítulos, estratégia_usada) ou (None, "").
    """
    # Estilos normais de corpo (não são títulos)
    NORMAL_STYLES = {
        "normal", "default paragraph style", "body text", "body text indent",
        "no spacing", "list paragraph", "quote", "intense quote",
        "footer", "header", "footnote text", "caption", "corpo",
    }
    HEADING_PREFIXES = ("heading", "título", "titulo", "title", "chapter", "capítulo")

    paras = _docx_collect_paragraphs(file_bytes)
    if not paras:
        return None, ""

    def _build(marker_fn, max_chapters: int = 300) -> list[dict] | None:
        chapters: list[dict] = []
        cur_title: str | None = None
        cur_lines: list[str] = []
        found = 0
        for p in paras:
            if marker_fn(p):
                found += 1
                content = "\n\n".join(cur_lines).strip()
                if content:
                    chapters.append({
                        "title": cur_title or "Prefácio / Introdução",
                        "content": content,
                    })
                cur_title = p["text"]
                cur_lines = []
            else:
                cur_lines.append(p["text"])
        content = "\n\n".join(cur_lines).strip()
        if content:
            chapters.append({"title": cur_title or "Sem título", "content": content})
        if found < 2 or found > max_chapters:
            return None
        # Sanidade: capítulos muito curtos indicam falsos positivos
        avg_words = sum(len(c["content"].split()) for c in chapters) / max(len(chapters), 1)
        if avg_words < 150:
            return None
        return chapters

    # ── Estratégia 1: estilos padrão ──────────────────────────────────────────
    result = _build(lambda p: any(p["style_lower"].startswith(h) for h in HEADING_PREFIXES))
    if result:
        return result, "estilos padrão Word (Heading/Title)"

    # ── Estratégia 2: estilo customizado frequente em parágrafos curtos ───────
    from collections import Counter
    style_counts: Counter = Counter(p["style"] for p in paras)
    dominant = style_counts.most_common(1)[0][0].lower() if style_counts else ""

    short_style: Counter = Counter(
        p["style"] for p in paras
        if p["is_short"] and p["style_lower"] not in NORMAL_STYLES and p["style_lower"] != dominant
    )
    candidates = [(s, c) for s, c in short_style.items() if c >= 2]
    if candidates:
        best = max(candidates, key=lambda x: x[1])[0]
        result = _build(lambda p, s=best: p["style"] == s)
        if result:
            return result, f"estilo customizado «{best}»"

    # Estratégia 3 (negrito) removida: documento com tema bold-by-default
    # causa falso positivo em todos os parágrafos. O fallback regex é mais confiável.

    return None, ""


def extract_pdf(file_bytes: bytes) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                txt = page.extract_text()
                if txt:
                    text_parts.append(txt)
        return "\n\n".join(text_parts)
    except ImportError:
        st.error("Biblioteca `pdfplumber` não encontrada. Execute: `pip install pdfplumber`")
        return ""
    except Exception as e:
        st.error(f"Erro ao ler PDF: {e}")
        return ""


def extract_text(uploaded_file) -> str:
    """Extrai texto de qualquer formato suportado."""
    raw = uploaded_file.read()
    name = uploaded_file.name.lower()
    if name.endswith(".docx"):
        return extract_docx(raw)
    elif name.endswith(".pdf"):
        return extract_pdf(raw)
    else:
        return extract_txt(raw)


# ─── Divisão automática em capítulos ─────────────────────────────────────────

CHAPTER_PATTERNS = [
    # "Capítulo 1", "Capítulo I", "CAPÍTULO 1", "Chapter 1"
    (r"^(?:cap[íi]tulo|chapter|cap\.?)\s+[\dIVXLCivxlc]+[.:\s-]*(.*)$", re.IGNORECASE),
    # "1- Título", "1. Título", "1: Título" — aceita todas as variantes de traço/ponto
    # \u2010=‐ \u2011=‑ \u2012=‒ \u2013=– \u2014=— \u2212=− (vários traços Unicode)
    (r"^\s*(\d{1,3})\s*[\-–—\.\:\u2010\u2011\u2012\u2013\u2014\u2212]+\s*(.{1,80})\s*$", re.IGNORECASE),
    # "## Título" (markdown h2 ou h1)
    (r"^#{1,3}\s+(.+)$", re.IGNORECASE),
    # Linha toda em MAIÚSCULAS — compilado SEM IGNORECASE para não capturar frases normais
    (r"^([A-ZÁÉÍÓÚÂÊÔÀÇÜÑ][A-ZÁÉÍÓÚÂÊÔÀÇÜÑ\s\d]{3,59})$", 0),
]


def split_into_chapters(text: str) -> list[dict]:
    """
    Tenta dividir o texto em capítulos detectando padrões de título.
    Retorna lista de dicts com 'title' e 'content'.
    """
    lines = text.splitlines()
    chapters = []
    current_title = None
    current_lines = []

    compiled = [re.compile(p, flags) for p, flags in CHAPTER_PATTERNS]

    def is_chapter_heading(line: str) -> str | None:
        stripped = line.strip()
        if not stripped or len(stripped) > 120:
            return None
        for pat in compiled:
            m = pat.match(stripped)
            if m:
                # Retorna o texto limpo do título
                groups = [g for g in m.groups() if g]
                return " ".join(groups).strip() if groups else stripped
        return None

    for line in lines:
        heading = is_chapter_heading(line)
        if heading:
            # Salva capítulo anterior
            if current_title is not None and "\n".join(current_lines).strip():
                chapters.append({
                    "title": current_title,
                    "content": "\n".join(current_lines).strip(),
                })
            elif current_lines and "\n".join(current_lines).strip() and current_title is None:
                # Texto antes do primeiro capítulo
                chapters.append({
                    "title": "Prefácio / Introdução",
                    "content": "\n".join(current_lines).strip(),
                })
            current_title = heading
            current_lines = []
        else:
            current_lines.append(line)

    # Último capítulo
    if current_lines and "\n".join(current_lines).strip():
        chapters.append({
            "title": current_title or "Capítulo 1",
            "content": "\n".join(current_lines).strip(),
        })

    return chapters if chapters else [{"title": "Capítulo 1", "content": text.strip()}]


# ─── Interface ────────────────────────────────────────────────────────────────

tab_single, tab_multi, tab_book, tab_free = st.tabs(
    ["📄 Um Capítulo", "📂 Vários Arquivos", "📚 Livro Inteiro (dividir automaticamente)", "✏️ Texto Livre"]
)

# ── Tab 1: Um arquivo = um capítulo ──────────────────────────────────────────
with tab_single:
    st.subheader("Importar como um único capítulo")
    st.caption("O conteúdo do arquivo será importado como um novo capítulo.")

    file_single = st.file_uploader(
        "Selecione o arquivo:",
        type=["txt", "md", "docx", "pdf"],
        key="upload_single",
    )

    if file_single:
        text = extract_text(file_single)
        if text:
            word_count = len(text.split())
            st.success(f"Arquivo lido: **{file_single.name}** · {word_count:,} palavras")

            chapter_name = st.text_input(
                "Nome do capítulo:",
                value=Path(file_single.name).stem.replace("_", " ").replace("-", " ").title(),
                key="single_chapter_name",
            )

            with st.expander("Pré-visualizar texto extraído", expanded=False):
                st.text_area("", value=text[:3000] + ("..." if len(text) > 3000 else ""), height=300, disabled=True)

            if st.button("✅ Importar como capítulo", type="primary", key="import_single"):
                new_ch = {
                    "id": str(uuid.uuid4()),
                    "title": chapter_name or file_single.name,
                    "content": text,
                    "notes": f"Importado de: {file_single.name}",
                }
                save_chapter(new_ch)
                st.success(f"Capítulo **{chapter_name}** importado com sucesso!")
                st.balloons()

# ── Tab 2: Múltiplos arquivos ─────────────────────────────────────────────────
with tab_multi:
    st.subheader("Importar vários arquivos como capítulos separados")
    st.caption("Cada arquivo se tornará um capítulo independente.")

    files_multi = st.file_uploader(
        "Selecione os arquivos (pode selecionar vários):",
        type=["txt", "md", "docx", "pdf"],
        accept_multiple_files=True,
        key="upload_multi",
    )

    if files_multi:
        st.markdown(f"**{len(files_multi)} arquivo(s) selecionado(s)**")

        previews = []
        for f in files_multi:
            text = extract_text(f)
            wc = len(text.split())
            default_name = Path(f.name).stem.replace("_", " ").replace("-", " ").title()
            previews.append({"file": f, "text": text, "name": default_name, "words": wc})

        # Edição de nomes antes de importar
        st.markdown("**Nomeie cada capítulo:**")
        names = []
        for i, p in enumerate(previews):
            col1, col2 = st.columns([3, 1])
            with col1:
                name = st.text_input(
                    f"Capítulo {i+1}:",
                    value=p["name"],
                    key=f"multi_name_{i}",
                )
                names.append(name)
            with col2:
                st.markdown(f"<br><small>{p['words']:,} palavras</small>", unsafe_allow_html=True)

        if st.button(f"✅ Importar {len(files_multi)} capítulo(s)", type="primary", key="import_multi"):
            for i, (p, name) in enumerate(zip(previews, names)):
                new_ch = {
                    "id": str(uuid.uuid4()),
                    "title": name or f"Capítulo {i+1}",
                    "content": p["text"],
                    "notes": f"Importado de: {p['file'].name}",
                }
                save_chapter(new_ch)
            st.success(f"**{len(files_multi)} capítulos** importados com sucesso!")
            st.balloons()

# ── Tab 3: Livro inteiro (divisão automática) ─────────────────────────────────
with tab_book:
    st.subheader("Importar livro inteiro e dividir em capítulos")
    st.markdown(
        "O app detecta automaticamente os títulos de capítulos no arquivo e os divide. "
        "Você pode revisar e ajustar antes de salvar."
    )

    file_book = st.file_uploader(
        "Selecione o arquivo do livro:",
        type=["txt", "md", "docx", "pdf"],
        key="upload_book",
    )

    if file_book:
        raw_bytes = file_book.read()
        _fname = file_book.name.lower()
        if _fname.endswith(".docx"):
            full_text = extract_docx(raw_bytes)
        elif _fname.endswith(".pdf"):
            full_text = extract_pdf(raw_bytes)
        else:
            full_text = extract_txt(raw_bytes)

        if full_text:
            total_words = len(full_text.split())
            st.success(f"**{file_book.name}** · {total_words:,} palavras no total")

            # Opções de divisão
            split_mode = st.radio(
                "Como dividir o livro:",
                [
                    "Automático (detectar títulos de capítulos)",
                    "Importar como um único capítulo",
                    "Dividir manualmente (por número de palavras)",
                ],
                key="book_split_mode",
            )

            detected_chapters = []

            if split_mode.startswith("Automático"):
                # Para .docx, tenta primeiro via análise do arquivo Word (mais preciso)
                docx_chaps, docx_strategy = None, ""
                if _fname.endswith(".docx"):
                    docx_chaps, docx_strategy = split_docx_into_chapters(raw_bytes)

                if docx_chaps:
                    detected_chapters = docx_chaps
                    st.info(
                        f"**{len(detected_chapters)} capítulo(s) detectado(s)** "
                        f"via {docx_strategy}. "
                        "Revise abaixo e ajuste os nomes se necessário."
                    )
                else:
                    detected_chapters = split_into_chapters(full_text)
                    st.info(
                        f"**{len(detected_chapters)} capítulo(s) detectado(s)** "
                        "via padrão de texto. "
                        "Revise abaixo e ajuste os nomes se necessário."
                    )

                # ── Debug: estilos encontrados no .docx ───────────────────────
                if _fname.endswith(".docx"):
                    with st.expander("🔍 Debug: estilos e parágrafos curtos do Word", expanded=not docx_chaps):
                        paras_meta = _docx_collect_paragraphs(raw_bytes)
                        if paras_meta:
                            from collections import Counter
                            style_counts = Counter(p["style"] for p in paras_meta)
                            st.markdown("**Estilos encontrados** (nome → nº de parágrafos):")
                            for s, c in style_counts.most_common():
                                st.caption(f"`{s}` → {c} parág.")
                            st.markdown("**Parágrafos curtos (≤80 chars) — caracteres reais:**")
                            st.caption("Negrito=✓ mostra parágrafos marcados como negrito.")
                            for p in paras_meta:
                                if p["is_short"]:
                                    bold_mark = "✓" if p["is_bold"] else "✗"
                                    st.caption(f"bold={bold_mark} `{repr(p['text'])}`")

            elif split_mode.startswith("Importar como"):
                book_name = st.text_input("Nome:", value=Path(file_book.name).stem.replace("_", " ").title())
                detected_chapters = [{"title": book_name, "content": full_text.strip()}]

            else:
                words_per_chapter = st.number_input(
                    "Palavras por capítulo:",
                    min_value=500,
                    max_value=20000,
                    value=3000,
                    step=500,
                    key="book_words_per_chapter",
                )
                words = full_text.split()
                chunks = [words[i:i+words_per_chapter] for i in range(0, len(words), words_per_chapter)]
                detected_chapters = [
                    {"title": f"Capítulo {i+1}", "content": " ".join(chunk)}
                    for i, chunk in enumerate(chunks)
                ]
                st.info(f"Será dividido em **{len(detected_chapters)} partes** de ~{words_per_chapter} palavras.")

            # ── Revisão dos capítulos detectados ──────────────────────────────
            if detected_chapters:
                st.markdown(f"### {len(detected_chapters)} capítulo(s) para importar")

                edited_chapters = []
                for i, ch in enumerate(detected_chapters):
                    with st.expander(
                        f"{'✏️' if i < 9 else '📄'} {i+1}. {ch['title']} "
                        f"({len(ch['content'].split()):,} palavras)",
                        expanded=False,
                    ):
                        new_title = st.text_input(
                            "Título:",
                            value=ch["title"],
                            key=f"book_title_{i}",
                        )
                        preview = st.text_area(
                            "Prévia (edite se necessário):",
                            value=ch["content"][:1500],
                            height=150,
                            key=f"book_preview_{i}",
                        )
                        # Guarda o texto completo (a prévia editável é só para visualização)
                        edited_chapters.append({
                            "title": new_title,
                            "content": ch["content"],  # Mantém o texto completo
                        })

                # Opção de substituir ou adicionar
                existing = get_chapters()
                if existing:
                    import_mode = st.radio(
                        f"Você já tem **{len(existing)} capítulo(s)**. O que fazer?",
                        ["Adicionar aos capítulos existentes", "Substituir todos os capítulos"],
                        key="book_import_mode",
                    )
                else:
                    import_mode = "Adicionar aos capítulos existentes"

                col_import, col_cancel = st.columns([2, 1])
                with col_import:
                    if st.button(
                        f"✅ Importar {len(detected_chapters)} capítulo(s)",
                        type="primary",
                        use_container_width=True,
                        key="import_book",
                    ):
                        if import_mode == "Substituir todos os capítulos":
                            book = load_book()
                            book["chapters"] = []
                            save_book(book)

                        for ch in (edited_chapters if edited_chapters else detected_chapters):
                            new_ch = {
                                "id": str(uuid.uuid4()),
                                "title": ch["title"],
                                "content": ch["content"],
                                "notes": f"Importado de: {file_book.name}",
                            }
                            save_chapter(new_ch)

                        total_imported = len(detected_chapters)
                        st.success(f"**{total_imported} capítulos** importados com sucesso!")
                        st.balloons()

# ── Tab 4: Texto Livre ────────────────────────────────────────────────────────
with tab_free:
    st.subheader("Importar texto colado diretamente")
    st.caption("Cole qualquer texto abaixo — pode ser um capítulo único ou um livro inteiro para dividir automaticamente.")

    free_text = st.text_area(
        "Cole o texto aqui:",
        height=350,
        placeholder="Cole aqui o conteúdo que deseja importar...",
        key="free_text_input",
    )

    if free_text.strip():
        word_count = len(free_text.split())
        st.caption(f"{word_count:,} palavras · {len(free_text):,} caracteres")

        free_mode = st.radio(
            "Como importar:",
            [
                "Como um único capítulo",
                "Dividir automaticamente em capítulos (detectar títulos)",
                "Dividir por número de palavras",
            ],
            key="free_import_mode",
        )

        # ── Modo: capítulo único ──────────────────────────────────────────────
        if free_mode == "Como um único capítulo":
            existing_chapters = get_chapters()

            free_action = st.radio(
                "Destino:",
                ["Criar novo capítulo", "Atualizar capítulo existente"],
                horizontal=True,
                key="free_action",
            )

            if free_action == "Criar novo capítulo":
                free_chapter_name = st.text_input(
                    "Nome do novo capítulo:",
                    value="Capítulo Importado",
                    key="free_chapter_name",
                )
                if st.button("✅ Criar capítulo", type="primary", key="import_free_single"):
                    new_ch = {
                        "id": str(uuid.uuid4()),
                        "title": free_chapter_name or "Capítulo Importado",
                        "content": free_text.strip(),
                        "notes": "Importado via texto livre",
                    }
                    save_chapter(new_ch)
                    st.success(f"Capítulo **{free_chapter_name}** criado com sucesso!")
                    st.balloons()

            else:  # Atualizar capítulo existente
                if not existing_chapters:
                    st.warning("Nenhum capítulo existente. Crie um primeiro no Editor.")
                else:
                    chapter_map = {ch["title"]: ch for ch in existing_chapters}
                    target_title = st.selectbox(
                        "Selecione o capítulo a atualizar:",
                        list(chapter_map.keys()),
                        key="free_update_target",
                    )
                    target_ch = chapter_map[target_title]
                    current_words = len(target_ch.get("content", "").split())
                    new_words = len(free_text.split())

                    col_cur, col_new = st.columns(2)
                    with col_cur:
                        st.metric("Palavras atuais", f"{current_words:,}")
                    with col_new:
                        delta = new_words - current_words
                        st.metric(
                            "Palavras novas",
                            f"{new_words:,}",
                            delta=f"{'+' if delta >= 0 else ''}{delta:,}",
                        )

                    with st.expander("Ver conteúdo atual do capítulo", expanded=False):
                        st.text(
                            target_ch.get("content", "")[:1500]
                            + ("..." if current_words > 200 else "")
                        )

                    st.warning(
                        f"O conteúdo atual de **{target_title}** será **substituído** pelo texto colado. "
                        "Esta ação não pode ser desfeita."
                    )

                    if st.button(
                        f"✅ Atualizar «{target_title}»",
                        type="primary",
                        key="import_free_update",
                    ):
                        updated_ch = dict(target_ch)
                        updated_ch["content"] = free_text.strip()
                        save_chapter(updated_ch)
                        st.success(f"Capítulo **{target_title}** atualizado com sucesso!")
                        st.balloons()

        # ── Modo: divisão automática ──────────────────────────────────────────
        elif free_mode == "Dividir automaticamente em capítulos (detectar títulos)":
            detected = split_into_chapters(free_text)
            st.info(
                f"**{len(detected)} capítulo(s) detectado(s)** automaticamente. "
                "Revise e ajuste os nomes se necessário."
            )

            edited_free = []
            for i, ch in enumerate(detected):
                with st.expander(
                    f"{i+1}. {ch['title']} ({len(ch['content'].split()):,} palavras)",
                    expanded=False,
                ):
                    new_title = st.text_input(
                        "Título:",
                        value=ch["title"],
                        key=f"free_title_{i}",
                    )
                    st.text_area(
                        "Prévia:",
                        value=ch["content"][:1000] + ("..." if len(ch["content"]) > 1000 else ""),
                        height=120,
                        disabled=True,
                        key=f"free_preview_{i}",
                    )
                    edited_free.append({"title": new_title, "content": ch["content"]})

            existing = get_chapters()
            if existing:
                free_replace = st.radio(
                    f"Você já tem **{len(existing)} capítulo(s)**. O que fazer?",
                    ["Adicionar aos capítulos existentes", "Substituir todos os capítulos"],
                    key="free_replace_mode",
                )
            else:
                free_replace = "Adicionar aos capítulos existentes"

            if st.button(
                f"✅ Importar {len(detected)} capítulo(s)",
                type="primary",
                key="import_free_auto",
            ):
                if free_replace == "Substituir todos os capítulos":
                    book = load_book()
                    book["chapters"] = []
                    save_book(book)

                for ch in edited_free:
                    new_ch = {
                        "id": str(uuid.uuid4()),
                        "title": ch["title"],
                        "content": ch["content"],
                        "notes": "Importado via texto livre",
                    }
                    save_chapter(new_ch)
                st.success(f"**{len(detected)} capítulos** importados com sucesso!")
                st.balloons()

        # ── Modo: divisão por palavras ────────────────────────────────────────
        else:
            free_wpc = st.number_input(
                "Palavras por capítulo:",
                min_value=500,
                max_value=20000,
                value=3000,
                step=500,
                key="free_words_per_chapter",
            )
            words = free_text.split()
            chunks = [words[i:i+free_wpc] for i in range(0, len(words), free_wpc)]
            st.info(f"Será dividido em **{len(chunks)} parte(s)** de ~{free_wpc} palavras.")

            if st.button(
                f"✅ Importar {len(chunks)} parte(s)",
                type="primary",
                key="import_free_split",
            ):
                for i, chunk in enumerate(chunks):
                    new_ch = {
                        "id": str(uuid.uuid4()),
                        "title": f"Capítulo {i + 1}",
                        "content": " ".join(chunk),
                        "notes": "Importado via texto livre",
                    }
                    save_chapter(new_ch)
                st.success(f"**{len(chunks)} capítulos** importados com sucesso!")
                st.balloons()

    else:
        st.info("Cole o texto acima para começar.")


# ─── Rodapé: estado atual ─────────────────────────────────────────────────────
st.divider()
chapters = get_chapters()
total_words = sum(len(c.get("content", "").split()) for c in chapters)
st.markdown(
    f"<small style='color:#aaa'>Manuscrito atual: **{len(chapters)}** capítulos · "
    f"**{total_words:,}** palavras</small>",
    unsafe_allow_html=True,
)
